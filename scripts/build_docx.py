import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from data.fetch import ZONES  # single source of truth, not a hand-kept copy

# Hand-listing these drifted: the document claimed Mahabalipuram, Pamban,
# Tuticorin and Mallipattinam -- none of which the system has ever
# evaluated -- while omitting Point Calimere, Mandapam, Rameswaram and
# Thoothukudi, which it does. A submitted architecture document naming
# four zones that do not exist is the kind of detail a reviewer checks.
_ZONE_NAMES = [z["name"] for z in ZONES]
ZONE_SENTENCE = (
    f"{len(_ZONE_NAMES)} coastal zones: "
    + ", ".join(_ZONE_NAMES[:-1])
    + f", and {_ZONE_NAMES[-1]}"
)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout(doc, title, text, bg_hex="F0F4F8", border_hex="1B365D"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"★ {title}\n")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    run_body = p.add_run(text)
    run_body.font.name = "Arial"
    run_body.font.size = Pt(10)
    run_body.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Add empty paragraph after table for spacing
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def format_run(run, font_name="Arial", size_pt=10.5, color_rgb=(0x33, 0x33, 0x33), bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.bold = bold
    run.italic = italic

def add_heading_1(doc, text):
    h = doc.add_heading(level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    format_run(run, font_name="Arial", size_pt=16, color_rgb=(0x1B, 0x36, 0x5D), bold=True)
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    format_run(run, font_name="Arial", size_pt=13, color_rgb=(0x0B, 0x6E, 0x5C), bold=True)
    return h

def add_heading_3(doc, text):
    h = doc.add_heading(level=3)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    format_run(run, font_name="Arial", size_pt=11.5, color_rgb=(0x22, 0x22, 0x22), bold=True)
    return h

def add_p(doc, text="", space_after=6, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        format_run(r_pre, font_name="Arial", size_pt=10.5, color_rgb=(0x1B, 0x36, 0x5D), bold=True)
    if text:
        r_text = p.add_run(text)
        format_run(r_text, font_name="Arial", size_pt=10.5, color_rgb=(0x33, 0x33, 0x33))
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        format_run(r_pre, font_name="Arial", size_pt=10.5, color_rgb=(0x1B, 0x36, 0x5D), bold=True)
    if text:
        r_text = p.add_run(text)
        format_run(r_text, font_name="Arial", size_pt=10.5, color_rgb=(0x33, 0x33, 0x33))
    return p

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8F9FA")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="0" w:color="CCCCCC"/><w:top w:val="single" w:sz="12" w:space="0" w:color="CCCCCC"/><w:right w:val="single" w:sz="12" w:space="0" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="CCCCCC"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def build_document():
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("ORCA System Architecture & Agentic Chatbot Workflow")
    format_run(run_title, font_name="Arial", size_pt=24, color_rgb=(0x1B, 0x36, 0x5D), bold=True)
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("Comprehensive Technical Documentation of Codebase, Deterministic Safety Core, and Fail-Closed AI Chatbot Layer")
    format_run(run_sub, font_name="Arial", size_pt=12, color_rgb=(0x0B, 0x6E, 0x5C), italic=True)
    
    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    r_div = p_div.add_run("―" * 55)
    format_run(r_div, font_name="Arial", size_pt=10, color_rgb=(0xCC, 0xCC, 0xCC))

    # SECTION 1: EXECUTIVE SUMMARY
    add_heading_1(doc, "1. Executive Summary & Core Engineering Philosophy")
    add_p(doc, "ORCA (Oceanic Risk & Condition Advisory) is a domain-specific marine safety reasoning system built for traditional and artisanal fishermen operating along the Tamil Nadu coastline (India). The platform evaluates " + ZONE_SENTENCE + ".")
    
    add_callout(doc, "The Core Architecture Principle: Zero-Trust AI & Boring Safety Core",
                "1. Deterministic Guarantee: Marine safety decisions (GO / DO NOT GO / SAFER ALTERNATIVE) are made exclusively by deterministic Python algorithms in policy.py and planner.py. Safety rules CANNOT be prompted out by an LLM.\n"
                "2. Fail-Closed AI Layer: The LLM agentic chatbot layer (orca/agentic.py) sits strictly around the safety engine. If the AI layer experiences any error (rate limit, key failure, timeout, invalid JSON), it degrades seamlessly to exact deterministic output byte-for-byte.\n"
                "3. Full Evidence Traceability: Every figure, wave height, wind speed, or temperature cited by the chatbot traces to a verified MarineObservation ID accessible via /evidence/{id}.")

    # SECTION 2: HIGH-LEVEL CODEBASE MAP
    add_heading_1(doc, "2. End-to-End Codebase Map & Directory Architecture")
    add_p(doc, "The repository is structured into distinct, modular layers adhering to strict boundaries:")

    # Table of files
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    set_cell_background(hdr_cells[0], "1B365D")
    set_cell_background(hdr_cells[1], "1B365D")
    set_cell_background(hdr_cells[2], "1B365D")
    
    hdr_titles = ["File / Directory", "Module Purpose", "Key Responsibilities & Safety Rules"]
    for i, t in enumerate(hdr_titles):
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(t)
        format_run(r, font_name="Arial", size_pt=10, color_rgb=(0xFF, 0xFF, 0xFF), bold=True)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)

    modules_info = [
        ("orca/agentic.py", "Agentic Chatbot Workflow", "Groq LLM integration for intent extraction, translation, and grounded phrasing. Never imports policy.py."),
        ("orca/planner.py", "Multi-Zone Planner", "Runs agents across 10 zones, evaluates alternatives, collects evidence IDs, builds Recommendations."),
        ("orca/policy.py", "Deterministic Safety Engine", "Zero network/LLM calls. Enforces hard denial precedence and risk override thresholds."),
        ("orca/agents.py", "Domain Evaluators", "5 independent agents: Satellite EO, Ocean State, Weather Risk, Wave Hazard, and Geofence/IMBL."),
        ("orca/memory.py", "Conversation Memory", "Sanitizes chat history into enum facts. NOTHING TYPED BY THE USER IS EVER STORED OR REPLAYED."),
        ("orca/api.py", "FastAPI Endpoints", "Exposes /ask, /evidence/{id}, /bathymetry, and /health. Delegates /ask to agentic layer."),
        ("orca/schema.py", "Data Models", "Pydantic & dataclass definitions for MarineObservation, Finding, and Recommendation."),
        ("data/fetch.py", "Ingestion & Caching", "Fetches multi-source ocean data (Open-Meteo, NOAA ERDDAP, MarineRegions IMBL API) to data/cache/."),
        ("web/index.html & three-viz.js", "Frontend UI & 3D Viz", "Vanilla CSS + Three.js visualizer. Displays honest AI-enhanced badges and native Tamil typography.")
    ]

    for row_idx, (fname, mname, desc) in enumerate(modules_info):
        row_cells = table.add_row().cells
        bg = "F9FBFD" if row_idx % 2 == 1 else "FFFFFF"
        for i, text in enumerate([fname, mname, desc]):
            set_cell_background(row_cells[i], bg)
            set_cell_margins(row_cells[i], top=80, bottom=80, left=120, right=120)
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            is_code = (i == 0)
            format_run(r, font_name="Consolas" if is_code else "Arial", size_pt=9.5 if is_code else 10, color_rgb=(0x1B, 0x36, 0x5D) if is_code else (0x33, 0x33, 0x33), bold=is_code)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

    # SECTION 3: DETERMINISTIC SAFETY CORE
    add_heading_1(doc, "3. The Deterministic Safety Core (orca/agents.py & policy.py)")
    add_p(doc, "Before an LLM ever touches a request, ORCA evaluates physical marine evidence through 5 domain agents and a strict decision tree:")
    
    add_bullet(doc, " Evaluates Chlorophyll-a (≥0.5 mg/m³) satellite imagery (NOAA VIIRS) and Sea Surface Temperature (27°C - 31°C) to identify fish aggregation zones.", "1. EO Satellite Agent (eo_satellite_agent):")
    add_bullet(doc, " Assesses thermal suitability and ocean current velocity.", "2. Ocean State Agent (ocean_state_agent):")
    add_bullet(doc, " Scales wind speed risk level up to 40 km/h.", "3. Weather Agent (weather_agent):")
    add_bullet(doc, " Evaluates significant wave height (m). Wave height > 2.5 m (Douglas Sea State 5 - Rough) triggers an unconditional HARD DENY.", "4. Wave Hazard Agent (hazard_agent):")
    add_bullet(doc, " Checks restricted marine zones (Gulf of Mannar Marine National Park MPA) and distance to the India-Sri Lanka Maritime Boundary (IMBL). IMBL < 2.0 km triggers an unconditional HARD DENY; < 5.0 km triggers warning; < 10.0 km triggers advisory.", "5. Geofence & IMBL Agent (geofence_agent):")

    add_heading_2(doc, "Decision Tree Rules in orca/policy.py")
    add_p(doc, "The decision engine resolves agent findings into a final zone action using 3 unyielding rules:")
    add_code_block(doc, 
"def resolve(findings: list[Finding]) -> Decision:\
    # Rule 1: Any hard denial (wave > 2.5m, MPA boundary, IMBL < 2km) wins unconditionally\
    if hard_denials:\
        return Decision(action='DO NOT GO', reason=primary.reason)\
    \
    # Rule 2: High risk (>= 0.6) contradicts fishing opportunity -> Safety Wins\
    if opportunity and danger:\
        return Decision(action='SAFER ALTERNATIVE', reason=danger[0].reason)\
    \
    # Rule 3: Acceptable conditions with no hazards -> GO\
    return Decision(action='GO', reason='No hazards found; conditions acceptable')")

    add_p(doc, "These three rules are the whole of orca/policy.py, which is frozen. They are not, however, the whole decision: orca/planner.py applies two corrections to each zone's Decision before it is accepted, one layer above the frozen module. Documenting the rules without them would overstate what policy.py alone guarantees.")
    add_bullet(doc, " A zone where no agent had any observation returns CANNOT ASSESS, not GO. Five neutral findings are not five clean bills of health, and \u201cNo hazards found; conditions acceptable\u201d from an empty evidence list is a confident answer to a question the system cannot answer. Deliberately not DO NOT GO: conflating \u201cI know it is dangerous\u201d with \u201cI do not know\u201d teaches users to discount the one verdict that must never be discounted.", "R-39 \u2014 no evidence is not safety:")
    add_bullet(doc, " Rule 2 above gates on opportunity AND danger, so a zone carrying a hazard with nothing suggesting go falls through to Rule 3 and returns GO. The trigger is inverted \u2014 suggests_go goes false when the water is cold or chlorophyll is cloud-masked \u2014 so the worse the fishing looks, the more likely the safety override is skipped. The planner returns SAFER ALTERNATIVE instead, naming the MOST SEVERE danger rather than the first in agent-registration order.", "R-59 \u2014 danger without opportunity:")

    # SECTION 4: AGENTIC CHATBOT WORKFLOW DETAILED BREAKDOWN
    add_heading_1(doc, "4. Detailed Agentic Chatbot Workflow (orca/agentic.py)")
    add_p(doc, "The chatbot in orca/agentic.py is implemented as a bounded, fixed-code-path WORKFLOW (following Anthropic's 'Building Effective Agents' design guidelines), NOT an autonomous agent picking its own tools. This ensures absolute predictability, speed, and safety.")

    add_heading_2(doc, "4.1 The Four-Tiered Zone Resolution Architecture")
    add_p(doc, "When a user submits a question ('Is it safe to fish near the southernmost tip of India today?'), ORCA resolves the maritime zone using a 4-tier hierarchy:")

    add_bullet(doc, "Zero Network / Zero Risk. The system checks if a known zone name (e.g. 'Nagapattinam') literally appears in the user query string. If matched, it uses this zone directly and completely bypasses the LLM.", "Tier 1: Literal Substring Match (_zone_by_substring) ― ")
    add_bullet(doc, "LLM Extraction. If Tier 1 finds no direct hit and GROQ_API_KEY is configured, the extraction model (gpt-oss-20b) maps landmarks/descriptions (e.g. 'southernmost tip of India' -> 'Kanyakumari') onto one of the 10 real zones in a closed Enum.", "Tier 2: LLM Intent & Landmark Extraction (extract_query_intent) ― ")
    add_bullet(doc, "Validated Conversation Memory. If the query names no place (e.g., 'What about tomorrow?'), the system resolves the zone from sanitized prior turns in memory.py.", "Tier 3: Prior Context Resolution (last_zone) ― ")
    add_bullet(doc, "Coordinate Fallback & Honest Caveat. If no zone can be inferred, the system defaults to the nearest geographic zone by latitude/longitude distance AND appends an explicit coverage note: 'You didn't name a place ORCA covers, so this is for Kanyakumari, the nearest of the 10 zones'.", "Tier 4: Geographic Fallback (build_recommendation) ― ")

    add_heading_2(doc, "4.2 Dual-Model Specialization Architecture")
    add_p(doc, "To maximize accuracy while respecting API rate limits (Groq free tier: 8,000 Tokens Per Minute), orca/agentic.py splits LLM duties across two specialized models:")

    # Model Table
    table2 = doc.add_table(rows=1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    for i, t in enumerate(["Stage", "Groq Model", "Config / Schema", "Purpose & Execution Strategy"]):
        set_cell_background(hdr2[i], "0B6E5C")
        set_cell_margins(hdr2[i], top=100, bottom=100, left=120, right=120)
        p = hdr2[i].paragraphs[0]
        r = p.add_run(t)
        format_run(r, font_name="Arial", size_pt=10, color_rgb=(0xFF, 0xFF, 0xFF), bold=True)

    model_rows = [
        ("1. Intent Extraction", "openai/gpt-oss-20b", "Temp = 0.0\nStrict JSON Schema", "Extracts zone_name, language ('en'/'ta'), intent ('verdict'/'data_lookup'), variable, time_frame, and on_topic. Fast & precise."),
        ("2. Grounded Composition", "openai/gpt-oss-120b", "Temp = 0.3\nStrict JSON Schema", "Phrases the deterministic decision in natural English or Tamil. Uses trimmed context (~200 tokens) to prevent rate limits.")
    ]

    for r_idx, (stg, mdl, cfg, prp) in enumerate(model_rows):
        rc = table2.add_row().cells
        bg = "F9FBFD" if r_idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate([stg, mdl, cfg, prp]):
            set_cell_background(rc[i], bg)
            set_cell_margins(rc[i], top=80, bottom=80, left=120, right=120)
            p = rc[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            format_run(r, font_name="Consolas" if i==1 else "Arial", size_pt=9.5 if i==1 else 10, color_rgb=(0x0B, 0x6E, 0x5C) if i==1 else (0x33, 0x33, 0x33))

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_after = Pt(12)

    add_heading_2(doc, "4.3 Trimmed Context Injection & Hallucination Immunization")
    add_p(doc, "Passing the entire raw recommendation dictionary (~3,200 tokens containing all 10 zone summaries, 5 raw agent findings, provenance URLs, and bathymetry grids) blew past Groq's 8,000 TPM limit after just two requests. Furthermore, handing excess raw JSON to an LLM increases hallucination risks.")
    add_p(doc, "orca/agentic.py introduces _composition_context() to trim context down to ~200 tokens:")

    add_code_block(doc,
"def _composition_context(recommendation: dict) -> dict:\
    # Minimal slice handed to the LLM composer\
    return {\
        'action': recommendation.get('action'),\
        'reason': recommendation.get('reason'),\
        'chosen_zone': (recommendation.get('chosen_zone') or {}).get('name'),\
        'evidence': [\
            {'id': e['id'], 'variable': e['variable'], 'value': e['value'], 'unit': e['unit']}\
            for e in recommendation.get('evidence', [])\
        ],\
    }")

    add_heading_2(doc, "4.4 Evidence Citation Validation & Fail-Closed Fallback")
    add_p(doc, "Even under strict JSON schema constraints, LLMs can occasionally cite non-existent evidence IDs. orca/agentic.py explicitly validates returned citations against real evidence IDs from planner.py:")
    add_code_block(doc,
"# Re-validate citations server-side:\
result['cited_evidence_ids'] = [\
    i for i in result.get('cited_evidence_ids', []) if i in real_evidence_ids\
]")

    # SECTION 5: CONVERSATION MEMORY ARCHITECTURE
    add_heading_1(doc, "5. Conversation Memory Architecture (orca/memory.py)")
    add_p(doc, "Multi-turn chat introduces two major security & reliability threats: Hallucination Compounding (carrying forward previous LLM errors) and Prompt Injection through Chat History ('ignore instructions...').")
    
    add_callout(doc, "The Absolute Memory Safety Rule (orca/memory.py)",
                "NOTHING THE USER TYPED IS EVER STORED OR REPLAYED. EVER.\n\n"
                "When a client sends chat history to /ask, memory.sanitize() instantly reduces every turn to at most 3 validated ENUM values: (zone_name, variable, time_frame). Free text is discarded entirely.\n\n"
                "- Immunization Against Hallucination Compounding: A bad prior turn leaves behind only a real zone name (e.g. 'Karaikal'), never bad prose. Every response is re-derived from fresh cached sensor data.\n"
                "- Immunization Against Prompt Injection: Injection strings cannot survive being cast into a closed Enum set.\n"
                "- Isolated Path: Sanitized turns reach ONLY the Intent Extraction step (to resolve missing subjects like 'what about tomorrow?'), and NEVER the Composition step.")

    # SECTION 6: FRONTEND & PROVENANCE TRACEABILITY
    add_heading_1(doc, "6. Frontend UX, Provenance Traceability & Testing")
    add_p(doc, "The user interface (web/index.html & three-viz.js) enforces visual honesty and complete data provenance:")
    
    add_bullet(doc, " The frontend displays an 'AI-enhanced' badge (#agentic-badge) ONLY when the LLM successfully executed. If the API key is missing or invalid, the badge is hidden, and exact deterministic text is displayed without misleading the user.", "Visual Honesty Badge: ")
    add_bullet(doc, " When Tamil ('ta') is detected, the UI sets lang='ta' on the response card, triggering native Tamil font styling and optimal line height.", "Multi-Language Typography: ")
    add_bullet(doc, " Every number shown in an answer is backed by an observation ID. Clicking an evidence item queries GET /evidence/{id}, returning full telemetry details (sensor source, timestamp, coordinates, confidence).", "Full Provenance Traceability: ")
    add_bullet(doc, " The entire agentic workflow is backed by 152 unit tests (pytest) and Playwright E2E suites (e2e/live.spec.js, e2e/agentic-exceptions.spec.js), proving that live HTTP 401s, 422s, 503s, and dead backend ports degrade gracefully without crashing.", "Comprehensive Automated Test Suite: ")

    # Save document
    out_path = "/home/anbu/26_class/ORCA/ORCA_Agentic_Chatbot_Architecture.docx"
    doc.save(out_path)
    print(f"Document successfully created at: {out_path}")

if __name__ == "__main__":
    build_document()
